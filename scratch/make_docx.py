import os
from PIL import Image, ImageDraw, ImageFont
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def draw_flowchart():
    # ── IMAGE CONFIGURATION ──
    w, h = 1200, 850
    img = Image.new("RGB", (w, h), "#F4F7F4") # Cool off-white background
    draw = ImageDraw.Draw(img)

    # Load font or fallback
    try:
        font_title = ImageFont.truetype("arial.ttf", 22)
        font_text = ImageFont.truetype("arial.ttf", 16)
        font_small = ImageFont.truetype("arial.ttf", 12)
        font_bold = ImageFont.truetype("arialbd.ttf", 15)
    except IOError:
        font_title = ImageFont.load_default()
        font_text = ImageFont.load_default()
        font_small = ImageFont.load_default()
        font_bold = ImageFont.load_default()

    # Colors
    color_border = "#2E7D32" # Forest Green
    color_orange = "#FF7043" # Coral Orange
    color_surface = "#FFFFFF" # Pure White Card
    color_text = "#111111"
    color_text_muted = "#555555"

    def draw_box(x_center, y_center, width, height, title, text="", border_color=color_border, is_diamond=False):
        left = x_center - width // 2
        top = y_center - height // 2
        right = x_center + width // 2
        bottom = y_center + height // 2

        if is_diamond:
            points = [
                (x_center, top),
                (right, y_center),
                (x_center, bottom),
                (left, y_center)
            ]
            draw.polygon(points, fill=color_surface, outline=border_color, width=3)
        else:
            # Rounded/standard rectangle outline and fill
            draw.rectangle([left, top, right, bottom], fill=color_surface, outline=border_color, width=3)

        # Draw Title
        title_w = draw.textlength(title, font=font_bold)
        draw.text((x_center - title_w // 2, y_center - (10 if text else 5)), title, fill=color_text, font=font_bold)
        
        # Draw Text
        if text:
            text_w = draw.textlength(text, font=font_small)
            draw.text((x_center - text_w // 2, y_center + 12), text, fill=color_text_muted, font=font_small)

    def draw_arrow(start, end, label=""):
        # Draw line
        draw.line([start, end], fill=color_border, width=3)
        
        # Draw arrowhead
        x1, y1 = start
        x2, y2 = end
        
        # Determine direction
        if x1 == x2: # Vertical arrow
            dy = 1 if y2 > y1 else -1
            arrow_pts = [(x2 - 8, y2 - dy * 12), (x2 + 8, y2 - dy * 12), (x2, y2)]
        elif y1 == y2: # Horizontal arrow
            dx = 1 if x2 > x1 else -1
            arrow_pts = [(x2 - dx * 12, y2 - 8), (x2 - dx * 12, y2 + 8), (x2, y2)]
        else: # Generic simple arrow
            arrow_pts = [(x2 - 8, y2 - 8), (x2 + 8, y2 - 8), (x2, y2)]
            
        draw.polygon(arrow_pts, fill=color_border)

        # Draw label if present
        if label:
            lx = (x1 + x2) // 2
            ly = (y1 + y2) // 2
            draw.text((lx + 10, ly - 10), label, fill=color_orange, font=font_bold)

    # ── DRAW FLOWCHART NODES ──
    # 1. Inputs Block
    draw_box(600, 70, 420, 80, "Grid Inputs", "Solar, Wind, Fossil Grid | Hospital, Industrial, Residential")
    
    # 2. Decision Diamond
    draw_box(600, 220, 240, 100, "Is Optimized?", "st.session_state.is_optimized", border_color=color_orange, is_diamond=True)

    # 3. No Branch - Base calculations
    draw_box(300, 390, 280, 70, "compute_base_efficiency", "Calculates initial baseline efficiency")
    draw_box(300, 520, 280, 70, "generate_unoptimized_distribution", "Chaotic, proportional fossil-heavy flows")

    # 4. Yes Branch - QAOA loop
    draw_box(900, 390, 280, 70, "run_qaoa_routing", "Initializes weights & controls training")
    draw_box(900, 520, 280, 70, "qaoa_circuit & qubo_cost", "4-Qubit simulator + Penalty calculations")

    # 5. Build Sankey
    draw_box(600, 670, 320, 70, "build_sankey(df, is_optimized)", "Generates flow mapping & node links")

    # 6. Render
    draw_box(600, 780, 320, 60, "Render Plotly Sankey View", "Interactive terminal output in browser")

    # ── DRAW FLOWCHART ARROWS ──
    # Inputs -> Decision
    draw_arrow((600, 110), (600, 170))

    # Decision -> No Branch (Left)
    draw.line([(600, 270), (300, 270)], fill=color_border, width=3)
    draw_arrow((300, 270), (300, 355), "NO")

    # Decision -> Yes Branch (Right)
    draw.line([(600, 270), (900, 270)], fill=color_border, width=3)
    draw_arrow((900, 270), (900, 355), "YES")

    # Left flow
    draw_arrow((300, 425), (300, 485))
    
    # Right flow
    draw_arrow((900, 425), (900, 485))
    
    # Loop back for QAOA loop training
    draw.arc([900 - 160, 390, 900 + 160, 520 + 35], start=270, end=90, fill=color_orange, width=3)
    draw.polygon([(1048, 455 - 6), (1048 + 12, 455 + 4), (1048 - 12, 455 + 4)], fill=color_orange)
    draw.text((1060, 445), "Gradient Descent (25 steps)", fill=color_orange, font=font_small)

    # Join Left flow to Sankey
    draw.line([(300, 555), (300, 615), (600, 615)], fill=color_border, width=3)
    draw_arrow((600, 615), (600, 635))

    # Join Right flow to Sankey
    draw.line([(900, 555), (900, 615), (600, 615)], fill=color_border, width=3)

    # Sankey -> Render
    draw_arrow((600, 705), (600, 750))

    # Save image
    img_path = "flowchart.png"
    img.save(img_path, "PNG", dpi=(300, 300))
    print("Flowchart image drawn and saved to", img_path)
    return img_path

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def create_word_document(image_path):
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style Configuration
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x11, 0x11, 0x11) # `#111111` (Black)

    # Theme Colors
    clr_primary = RGBColor(0x2E, 0x7D, 0x32) # `#2E7D32` (Forest Green)
    clr_secondary = RGBColor(0x55, 0x55, 0x55) # `#555555` (Muted Gray)

    # Helper for adding styled headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = clr_primary
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = clr_primary
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        return p

    def add_bullet(bold_prefix, text_content):
        p = doc.add_paragraph(style='List Bullet')
        run_bold = p.add_run(bold_prefix + ": ")
        run_bold.bold = True
        run_bold.font.color.rgb = clr_primary
        p.add_run(text_content)
        p.paragraph_format.space_after = Pt(3)

    # ── DOCUMENT HEADER ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("EcoGrid.AI — Technical Stack & Core Logic Specification")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = clr_primary
    title_p.paragraph_format.space_after = Pt(2)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Enterprise-Grade Quantum-Optimized Renewable Energy Distribution Framework")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = clr_secondary
    subtitle_p.paragraph_format.space_after = Pt(24)

    # Add a thin line
    p_line = doc.add_paragraph()
    p_line_run = p_line.add_run("━" * 57)
    p_line_run.font.color.rgb = clr_primary
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_after = Pt(20)

    # ── SECTION 1: ARCHITECTURAL OVERVIEW ──
    add_heading_1("1. Architectural Overview & Constraints")
    p_arch = doc.add_paragraph(
        "EcoGrid.AI is a cutting-edge quantum energy optimization simulator built for hackathons under high operational constraints. "
        "The complete system runs under a flat, single-file monolithic structure (app.py) designed for high performance, local execution, and low latency."
    )
    p_arch.paragraph_format.space_after = Pt(10)

    add_bullet("Offline Execution Model", "The entire quantum simulation and resource allocation run locally on host CPUs. The engine relies on no third-party APIs or remote cloud-quantum services.")
    add_bullet("Single-File Delivery", "Logic, visual formatting, and numerical pipelines reside strictly inside app.py to bypass pathing and multi-repository linkage errors.")
    add_bullet("Reactive State Management", "Leverages Streamlit's session states to coordinate sliders, crisis preset scenarios, and optimal matrices seamlessly.")

    # ── SECTION 2: SYSTEM DEPENDENCIES ──
    add_heading_1("2. Ecosystem Dependencies & Tech Stack")
    doc.add_paragraph(
        "The environment uses standard, robust open-source scientific packages to simulate quantum states and map supply-to-demand flow relationships:"
    )
    
    add_bullet("Streamlit (>= 1.30.0)", "Provides state holding, UI design widgets, and layouts. The interface leverages custom inline HTML & CSS for high-fidelity rendering.")
    add_bullet("PennyLane (>= 0.34.0)", "Main variational quantum simulation framework. Used to establish the simulator device ('default.qubit') and QAOA circuit layout.")
    add_bullet("Plotly (>= 5.18.0)", "Renders real-time, interactive Sankey Diagrams showing distribution flow from generators to residential/industrial zones.")
    add_bullet("Pandas (>= 2.0.0)", "Structures the raw routing matrix rows (Source, Target, MW Flow) before and after optimal optimization loops.")

    # ── SECTION 3: SYSTEM ARCHITECTURE FLOWCHART ──
    add_heading_1("3. System Architecture & Flowchart")
    doc.add_paragraph(
        "The following diagram illustrates the functional logic pipeline. It traces how user slider inputs trigger the reactive flow—shifting from baseline unoptimized routing models into the variational PennyLane optimization circuit."
    )

    # Add Diagram Image
    doc.add_picture(image_path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run("Figure 1.0: EcoGrid.AI Data Processing and Variational QAOA Optimization Circuit Pipeline.")
    caption_run.font.size = Pt(9.5)
    caption_run.italic = True
    caption_run.font.color.rgb = clr_secondary
    caption_p.paragraph_format.space_after = Pt(20)

    # ── SECTION 4: CORE FUNCTION IMPLEMENTATION ──
    add_heading_1("4. Core Logical & Quantum Functions")
    doc.add_paragraph(
        "The core code is organized into modular logical and mathematical layers. Below is a structured analysis of the primary functions powering the application:"
    )

    # Create table for functions
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    
    # Format Headers
    hdr_cells = table.rows[0].cells
    headers = ["Function Name", "Logical Role", "Mathematical Operation / Detail"]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "2E7D32")
        # Text properties for header
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    # Function Details Data
    func_data = [
        ("compute_base_efficiency", 
         "Calculates the baseline efficiency score of the grid without QAOA optimization.", 
         "Compares supply vs demand. Penalizes fossil dependence: Penalty = (Fossil / Total) * 20. Clamps output between 10% and 70%."),
        
        ("generate_unoptimized_distribution", 
         "Creates a standard proportional supply-to-demand flow representing unoptimized conditions.", 
         "Spreads Solar, Wind, and Fossil fuel supply across sectors with heavy, unoptimized fossil bias."),
        
        ("qaoa_circuit", 
         "Defines the 4-qubit parameterized variational circuit on the PennyLane default.qubit simulator.", 
         "Prepares equal superposition with Hadamards, applies variational RX/RZ layer pairs, entangles via cyclic CNOT rings, and outputs 16-state probabilities."),
        
        ("qubo_cost", 
         "Cost function mapping optimization constraints and penalties onto quantum states.", 
         "Penalizes hospital shortages (+8.0), unnecessary fossil usage (+5.0), blackout risks (+4.0), and rewards renewable usage (-1.5 per active route)."),
        
        ("run_qaoa_routing", 
         "Initializes weights and optimizes parameters, then decodes the quantum state into real power flows.", 
         "Runs GradientDescentOptimizer for 25 steps, extracts the argmax probability state, and converts binary nodes (Q0-Q3) into exact MW flow targets."),
        
        ("build_sankey", 
         "Generates a structured Plotly Sankey object mapping flows from sources to demand sinks.", 
         "Translates DataFrame rows into visual nodes and paths, adapting styling colors based on whether optimization is active.")
    ]

    for name, role, math in func_data:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        # Bold function names
        for p in row_cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = clr_primary
        row_cells[1].text = role
        row_cells[2].text = math

    # Spacing after table
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Pt(12)

    # ── SECTION 5: QUANTUM LOGIC BREAKDOWN ──
    add_heading_1("5. In-Depth Variational Circuit Mapping")
    doc.add_paragraph(
        "At the core of the quantum engine is the mapping of grid-distribution constraints to the Hilbert space of four qubits. "
        "Because each qubit represents a specific binary decision (e.g., routing solar/wind energy to the hospital or residential grids), "
        "the circuit optimizes the combination of these decisions simultaneously."
    )
    
    p_qubo = doc.add_paragraph()
    p_qubo.add_run("The QUBO cost function mathematically translates to a Hamiltonian:").italic = True
    p_qubo.paragraph_format.space_before = Pt(6)
    p_qubo.paragraph_format.space_after = Pt(6)
    
    # Formula represented cleanly
    formula_p = doc.add_paragraph()
    formula_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula_run = formula_p.add_run(
        "H_cost = w_1 * P_hospital + w_2 * P_fossil_waste + w_3 * P_blackout - w_4 * R_green"
    )
    formula_run.bold = True
    formula_run.font.size = Pt(12)
    formula_run.font.color.rgb = clr_primary
    formula_p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(
        "During execution, the Gradient Descent loop updates the variational parameters to shift the probability amplitude "
        "toward the exact state that minimizes H_cost. In a stable system, the system converges on state combinations where "
        "critical systems are covered by green energy, fossil fuel backup usage is restricted only to crisis conditions, "
        "and carbon footprints are strictly minimized."
    )

    # Save Docx
    doc_path = "details.docx"
    doc.save(doc_path)
    print("Word document generated and saved to", doc_path)
    return doc_path

if __name__ == "__main__":
    img_path = draw_flowchart()
    doc_path = create_word_document(img_path)
    # Clean up temp image
    if os.path.exists(img_path):
        os.remove(img_path)
        print("Temporary image cleaned up.")
